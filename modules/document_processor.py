"""Safe multi-format ingestion with page/sheet evidence and optional OCR fallback."""
import os
from dataclasses import dataclass, asdict
from typing import Any, Dict, List
import PyPDF2
from docx import Document
import pandas as pd
import config

@dataclass
class DocumentEvidence:
    source: str
    location: str
    text: str
    kind: str = "text"
    def to_dict(self): return asdict(self)

class DocumentProcessor:
    def __init__(self): self.supported_formats = set(config.SUPPORTED_FORMATS)
    def validate_file(self, file_path):
        if not os.path.isfile(file_path): raise ValueError("Document does not exist.")
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.supported_formats: raise ValueError(f"Unsupported format: {ext}. Supported: {sorted(self.supported_formats)}")
        if os.path.getsize(file_path) > config.MAX_FILE_SIZE_MB * 1024 * 1024: raise ValueError(f"File exceeds {config.MAX_FILE_SIZE_MB} MB limit.")
    def extract(self, file_path):
        self.validate_file(file_path); ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf": text, evidence = self._pdf(file_path)
        elif ext == ".docx": text, evidence = self._docx(file_path)
        elif ext == ".txt": text, evidence = self._txt(file_path)
        else: text, evidence = self._xlsx(file_path)
        truncated = len(text) > config.MAX_DOCUMENT_CHARS
        return {"text": text[:config.MAX_DOCUMENT_CHARS], "evidence": evidence, "truncated": truncated,
                "source": os.path.basename(file_path), "characters": min(len(text), config.MAX_DOCUMENT_CHARS)}
    def extract_text(self, file_path): return self.extract(file_path)["text"]
    def _pdf(self, path):
        pages, evidence = [], []; ocr_pages = 0
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages, 1):
                page_text = (page.extract_text() or "").strip(); kind = "text"
                if not page_text and config.OCR_ENABLED:
                    page_text = self._ocr_page(path, i) or ""; kind = "ocr" if page_text else "text"
                    ocr_pages += int(bool(page_text))
                pages.append(f"[Page {i}]\n{page_text}")
                evidence.append(DocumentEvidence(os.path.basename(path), f"page {i}", page_text[:4000], kind).to_dict())
        if not pages: raise ValueError("PDF contains no readable pages.")
        return "\n\n".join(pages), evidence
    def _ocr_page(self, path, page_number):
        try:
            import fitz, pytesseract
            from PIL import Image
            doc = fitz.open(path); page = doc.load_page(page_number - 1)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            return pytesseract.image_to_string(image).strip()
        except Exception:
            return ""
    def _docx(self, path):
        doc = Document(path); blocks=[]; evidence=[]
        for i,p in enumerate(doc.paragraphs,1):
            if p.text.strip():
                blocks.append(p.text.strip()); evidence.append(DocumentEvidence(os.path.basename(path),f"paragraph {i}",p.text.strip()[:4000]).to_dict())
        for ti,table in enumerate(doc.tables,1):
            rows=[" | ".join(c.text.strip() for c in row.cells) for row in table.rows]
            text="\n".join(rows)
            if text: blocks.append(f"[Table {ti}]\n{text}"); evidence.append(DocumentEvidence(os.path.basename(path),f"table {ti}",text[:4000],"table").to_dict())
        return "\n".join(blocks), evidence
    def _txt(self,path):
        with open(path,"r",encoding="utf-8",errors="replace") as f: text=f.read()
        return text,[DocumentEvidence(os.path.basename(path),"text file",text[:4000]).to_dict()]
    def _xlsx(self,path):
        sheets=pd.read_excel(path,sheet_name=None); blocks=[]; evidence=[]
        for name,frame in sheets.items():
            text=frame.fillna("").to_string(index=False); blocks.append(f"[Sheet: {name}]\n{text}"); evidence.append(DocumentEvidence(os.path.basename(path),f"sheet {name}",text[:4000],"table").to_dict())
        return "\n\n".join(blocks),evidence
